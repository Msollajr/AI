from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x0C, 0x18, 0x36)

# ── Helper functions ─────────────────────────────────────────────────────────
def add_code_block(doc, text, font_size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table_row(table, cells_text, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(10)
        if bold:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return row

def add_diagram_block(doc, text, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)
    add_code_block(doc, text, font_size=8)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()

# ── UDSM Logo on title page ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
run.add_picture('frontend/udsm-logo.png', width=Inches(1.8))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('IS 365 — Practical Assignment')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x0C, 0x18, 0x36)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Full-Stack Pipeline for Deploying\na Self-Hosted LLM Application')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x0C, 0x18, 0x36)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('University Student Support LLM Application')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x6C, 0xAB, 0xDD)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Generated: {datetime.datetime.now().strftime("%B %d, %Y")}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction',
    '2. System Use Case',
    '3. Tools and Technologies Used',
    '4. System Architecture',
    '5. Implementation Steps',
    '   5.1 Task 1: Environment Setup',
    '   5.2 Task 2: Local LLM Setup',
    '   5.3 Task 3: Backend Development',
    '   5.4 Task 4: Frontend Development',
    '   5.5 Task 5: Testing Setup',
    '   5.6 Task 6: Prompt Engineering',
    '   5.7 Task 7: Error Handling',
    '   5.8 Task 8: Logging',
    '6. Testing and Results',
    '7. Challenges Encountered',
    '8. Bonus Features',
    '9. Industry Production Reflection',
    '10. Conclusion',
    'Appendix A: Project Structure',
    'Appendix B: Prompt Comparison',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'Modern artificial intelligence deployment extends far beyond writing an LLM inference script. '
    'In production environments, an AI system is an orchestration of multiple connected layers: '
    'virtual environments, local model servers, robust API backends, web frontends, automated test '
    'suites, structured logging, and thorough error handling.'
)
doc.add_paragraph(
    'This report documents the design, implementation, and deployment of a self-hosted University '
    'Student Support Assistant. By running a local Small Language Model (SLM) on an offline server, '
    'this architecture demonstrates how organizations can leverage generative AI while maintaining '
    'absolute data ownership, zero API latency costs, and complete privacy.'
)
doc.add_paragraph(
    'The system uses Ollama to serve the llama3.2:1b model locally, FastAPI as the backend API '
    'framework, a Tailwind CSS Single Page Application (SPA) as the frontend, and a Streamlit '
    'wrapper for deployment on an alternative port. The pipeline includes Retrieval-Augmented '
    'Generation (RAG) using 11 official university documents, semantic caching, conversation '
    'memory, confidence scoring, accuracy validation, and structured response formatting.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM USE CASE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('2. System Use Case', level=1)
doc.add_paragraph(
    'The primary use case is to assist students with standard administrative inquiries across '
    'eight core areas of university services:'
)

use_cases = [
    ('Course Registration: ', 'Timelines, portal navigation, adding/dropping classes.'),
    ('Examination Rules: ', 'Attendance rules, prohibited items, grading standards, and supplementary exams.'),
    ('Library Services: ', 'Opening hours, lending limits, OPAC search, online databases, and overdue fines.'),
    ('ICT Support: ', 'Campus Wi-Fi configuration, email login, and password resetting procedures.'),
    ('Hostel Application: ', 'On-campus accommodation allocation priority, portal application processes, and fees.'),
    ('Fee Payment: ', 'Payment channels, installment structures, and registration blocks for outstanding balances.'),
    ('Academic Calendar: ', 'Lecture weeks, mid-semester breaks, study leave, and exam periods.'),
    ('Student Conduct: ', 'Behavior codes, prohibited substances, dress codes, and disciplinary hearings.'),
]
for bold_part, text in use_cases:
    add_bullet(doc, text, bold_prefix=bold_part)

doc.add_paragraph(
    'To prevent the LLM from hallucinating administrative policies, we implemented a '
    'Retrieval-Augmented Generation (RAG) pipeline using a knowledge base of 11 official UDSM '
    'markdown documents alongside a keyword-matched FAQ database. The system also includes a '
    'FAQ Direct Answer engine that bypasses the LLM entirely for known questions, providing '
    'sub-second responses with guaranteed accuracy.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. TOOLS AND TECHNOLOGIES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Tools and Technologies Used', level=1)
doc.add_paragraph('The project stack was selected for lightweight local execution, performance, and standard industry practices:')

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr = table.rows[0]
for i, text in enumerate(['Component', 'Tool / Library', 'Justification']):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

tools_data = [
    ('Operating System', 'Windows', 'Host environment.'),
    ('Language & Environment', 'Python 3.13 & venv', 'Package isolation and standard scripting.'),
    ('Backend API Framework', 'FastAPI', 'Async endpoints, validation, Swagger UI.'),
    ('Web Server', 'Uvicorn', 'Lightning-fast ASGI server.'),
    ('Local LLM Server', 'Ollama', 'Industry-standard local model serving.'),
    ('Model', 'llama3.2:1b', '1B parameter model for low-latency local execution.'),
    ('Frontend', 'Tailwind CSS SPA + Streamlit', 'Production SPA + Streamlit wrapper on port 8501.'),
    ('Testing Suite', 'Pytest & TestClient', 'Unit testing with FastAPI TestClient.'),
]
for comp, tool, justification in tools_data:
    add_table_row(table, [comp, tool, justification])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. System Architecture', level=1)
doc.add_paragraph('The application runs as a modular, three-tier local architecture:')

arch_diagram = """                   ┌──────────────────────┐
                   │     Student User     │
                   └──────────────────────┘
                               │
                       User Input / Chat
                               ▼
                   ┌──────────────────────┐
                   │   Tailwind CSS SPA   │  (index.html)
                   │   Streamlit Wrapper  │  (app.py, port 8501)
                   └──────────────────────┘
                         │            ▲
                  POST  /ask        JSON
                     Request       Response
                         ▼            │
                   ┌──────────────────────────────────────────┐
                   │            FastAPI Backend               │
                   │  ┌─────────┐  ┌──────────┐  ┌─────────┐ │
                   │  │FAQ      │  │Semantic  │  │Convers- │ │
                   │  │Direct   │  │Cache     │  │ation    │ │
                   │  │Answer   │  │(30min    │  │Memory   │ │
                   │  │Engine   │  │TTL)      │  │(6turns) │ │
                   │  └─────────┘  └──────────┘  └─────────┘ │
                   │  ┌─────────┐  ┌──────────┐              │
                   │  │Accuracy │  │Confidence│              │
                   │  │Check    │  │Scoring   │              │
                   │  └─────────┘  └──────────┘              │
                   └──────────────────────────────────────────┘
                      │        ▲           │
            Lookup KB  │        │ Context   │ Log interactions
            + FAQ      ▼        │ Loaded    │ + ratings
                   ┌──────────────┐    ┌─────────────────┐
                   │  faq.json    │    │  logs/app.log   │
                   │  knowledge-  │    │  logs/feed.log  │
                   │  base/*.md   │    └─────────────────┘
                   └──────────────┘
                         │            ▲
                  POST  /api/generate │
                    with context      │
                         ▼            │
                   ┌──────────────────────┐
                   │   Ollama Local LLM   │  (llama3.2:1b)
                   └──────────────────────┘"""

add_diagram_block(doc, arch_diagram, 'Figure 1: Complete System Architecture')

doc.add_paragraph()
doc.add_heading('Data Flow', level=2)
doc.add_paragraph('The system processes each student query through the following pipeline:')

steps = [
    'The student submits a question via the SPA chat interface.',
    'FastAPI Backend receives the question and performs a cache lookup (SemanticCache — 30-min TTL, 0.92 similarity threshold).',
    'If not cached, the FAQ Direct Engine checks for a high-confidence match (>=85%). If found, returns the answer immediately without invoking the LLM.',
    'If no FAQ match, the Knowledge Base RAG engine retrieves relevant sections from 11 markdown documents using keyword overlap scoring.',
    'The Conversation Memory checks for repeated or context-dependent questions (last 6 turns, SequenceMatcher > 0.85).',
    'The backend builds an expert prompt with system instructions, retrieved context, and conversation history.',
    'The prompt is sent to Ollama, which processes it on the llama3.2:1b model.',
    'The backend runs an accuracy check (hallucination guardrails, minimum length threshold), calculates a confidence score (hybrid: FAQ score + RAG score + source count), and optionally checks the UDSM website for time-sensitive updates.',
    'The backend wraps the response in the structured 5-section format (Answer, Steps, Useful Links, Sources, Confidence).',
    'The interaction is logged to app.log with timestamp, question, answer, and metadata.',
    'The frontend displays the formatted response with color-coded confidence badges, source citations, and feedback buttons.',
]
for i, step in enumerate(steps, 1):
    add_bullet(doc, step, bold_prefix=f'Step {i}: ')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. IMPLEMENTATION STEPS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Implementation Steps', level=1)

# 5.1
doc.add_heading('5.1 Task 1: Environment Setup', level=2)
doc.add_paragraph('The development environment was prepared with the following steps:')
add_bullet(doc, 'Created Python virtual environment: python -m venv venv', 'Command: ')
add_bullet(doc, 'Activated the environment and installed all dependencies from requirements.txt.', 'Action: ')
doc.add_paragraph('The requirements.txt includes the following packages:')

reqs_table = doc.add_table(rows=1, cols=2)
reqs_table.style = 'Light Grid Accent 1'
reqs_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = reqs_table.rows[0]
for i, text in enumerate(['Package', 'Purpose']):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

packages = [
    ('fastapi', 'Backend API framework'),
    ('uvicorn', 'ASGI web server'),
    ('streamlit', 'Frontend wrapper'),
    ('requests', 'HTTP client for Ollama'),
    ('httpx', 'HTTP client (TestClient)'),
    ('pytest', 'Testing framework'),
    ('pydantic', 'Data validation'),
    ('python-multipart', 'Form data parsing'),
]
for pkg, purpose in packages:
    add_table_row(reqs_table, [pkg, purpose])

# 5.2
doc.add_heading('5.2 Task 2: Local LLM Setup', level=2)
doc.add_paragraph('Ollama was installed and configured to serve the llama3.2:1b model locally:')
add_bullet(doc, 'ollama pull llama3.2:1b — pulled the 1.3B parameter model (Q8_0 quantization).', 'Command: ')
add_bullet(doc, 'ollama run llama3.2:1b — verified the model responds to prompts.', 'Command: ')
add_bullet(doc, 'The Ollama REST API is available at http://localhost:11434/api/generate.', 'Verified: ')

# 5.3
doc.add_heading('5.3 Task 3: Backend Development', level=2)
doc.add_paragraph('The backend consists of three core files:')

doc.add_heading('config.py', level=3)
doc.add_paragraph(
    'Environment configuration module that loads host, port, Ollama URL, model name, and file paths. '
    'Uses environment variables with sensible defaults for local development.'
)

doc.add_heading('llm_client.py — Core LLM Client', level=3)
doc.add_paragraph('The LLMClient class implements the complete RAG pipeline with the following components:')

components = [
    ('ConversationMemory: ', 'Manages the last 6 conversation turns, detects repeated questions via SequenceMatcher (threshold > 0.85), and returns conversation context for prompt building.'),
    ('SemanticCache: ', '30-minute TTL cache with 0.92 similarity threshold and maximum 100 entries. Caches both FAQ direct answers and LLM-generated responses.'),
    ('FAQDirectAnswer: ', 'Keyword-matching engine that checks the FAQ database. Returns answer without invoking the LLM when confidence >= 0.85 (sub-second responses). Injects FAQ context into the prompt when confidence is between 0.60 and 0.85.'),
    ('retrieve_context(): ', 'Dual retrieval strategy: first checks FAQ keywords, then performs full-text search across 11 knowledge-base markdown documents using word-overlap scoring. Returns context string, category, sources list, and RAG score.'),
    ('construct_expert_prompt(): ', 'Builds a system prompt with the instruction to answer based only on provided UDSM information. Injects retrieved context and conversation history into the prompt.'),
    ('generate_response(): ', 'End-to-end pipeline method orchestrating: cache check → FAQ direct → context retrieval → prompt construction → LLM query → accuracy check → confidence calculation → web check → structured formatting → memory storage → cache update.'),
    ('_calculate_confidence(): ', 'Hybrid confidence scoring based on FAQ match score, RAG overlap score, and number of unique sources. Returns label (High/Medium/Low) and numeric score (0.0–1.0).'),
    ('_accuracy_check(): ', 'Hallucination detection that checks for refusal indicators ("I cannot provide", "unfortunately", "I\'m sorry") and minimum answer length (< 20 characters).'),
    ('_format_structured_response(): ', 'Wraps the raw LLM answer in the 5-section format: Answer, What You Should Do (extracted via regex action-verb matching), Useful Links (5 fixed UDSM URLs), Sources (from retrieval), and Confidence (label + percentage).'),
]
for bold_part, text in components:
    add_bullet(doc, text, bold_prefix=bold_part)

doc.add_heading('main.py — FastAPI Application', level=3)
doc.add_paragraph('The FastAPI application exposes four endpoints:')

endpoints_table = doc.add_table(rows=1, cols=3)
endpoints_table.style = 'Light Grid Accent 1'
endpoints_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = endpoints_table.rows[0]
for i, text in enumerate(['Method', 'Path', 'Description']):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

endpoints = [
    ('GET', '/', 'Serves the Tailwind CSS SPA (index.html)'),
    ('GET', '/health', 'System health check — validates Ollama connectivity and model presence'),
    ('POST', '/ask', 'Accepts student question, returns structured response with 5-section format'),
    ('POST', '/feedback', 'Logs user response rating (Good / Average / Poor)'),
]
for method, path, desc in endpoints:
    add_table_row(endpoints_table, [method, path, desc])

doc.add_paragraph()
doc.add_paragraph(
    'The /ask endpoint returns an AskResponse model with the following fields: answer, '
    'rag_context_used, category, sources (list), confidence_label (High/Medium/Low), '
    'confidence_score (float 0.0–1.0), faq_direct (bool), and timestamp. The endpoint '
    'includes comprehensive error handling: 400 for empty questions, 503 when Ollama is '
    'unreachable, and standard 422 for schema validation failures.'
)

# 5.4
doc.add_heading('5.4 Task 4: Frontend Development', level=2)
doc.add_paragraph('The frontend consists of two components:')

doc.add_heading('Tailwind CSS SPA (index.html)', level=3)
doc.add_paragraph('A fully responsive Single Page Application built with Tailwind CSS and Lucide icons:')
add_bullet(doc, 'Chat interface with message bubbles, welcome splash screen, and animated loading spinner.', 'Features: ')
add_bullet(doc, 'Color-coded confidence badges (High — green, Medium — amber, Low — red).', 'Confidence Display: ')
add_bullet(doc, 'Source citations displayed beneath each assistant message.', 'Source Display: ')
add_bullet(doc, '5-section response rendering (Answer, Steps, Useful Links, Sources, Confidence).', 'Structured Responses: ')
add_bullet(doc, 'Feedback buttons (Good / Average / Poor) on each assistant message.', 'Feedback: ')
add_bullet(doc, 'Dark mode toggle with persistence.', 'Dark Mode: ')
add_bullet(doc, 'Citations sidebar panel showing retrieved category and sources.', 'Citations Panel: ')
add_bullet(doc, 'Dashboard tab with mock student information cards (registration status, hostel allocation, fee status).', 'Dashboard: ')
add_bullet(doc, 'Mobile-first responsive design with sidebar overlay navigation on small screens, adaptive typography, and touch-friendly interaction targets.', 'Responsive: ')

doc.add_heading('Streamlit Wrapper (app.py)', level=3)
doc.add_paragraph(
    'A thin 46-line Streamlit application that embeds the SPA in a full-viewport iframe. This '
    'provides the required deployment on port 8501 while serving the identical SPA experience. '
    'The wrapper hides all Streamlit chrome elements and uses CSS media queries for proper '
    'mobile viewport handling (100dvh, min-width: 320px).'
)

# 5.5
doc.add_heading('5.5 Task 5: Testing Setup', level=2)
doc.add_paragraph(
    'A comprehensive test suite was created using Pytest with FastAPI\'s TestClient. The '
    'test file (tests/test_api.py) contains 17 individual tests organized into five test classes:'
)

test_classes = [
    ('TestHealthEndpoint (2 tests): ', 'Verifies the health endpoint returns HTTP 200 and contains all expected fields (status, llm_connected, timestamp).'),
    ('TestAskEndpoint (9 tests): ', 'Validates question rejection (empty, whitespace, missing), response structure (all 7 fields present), FAQ direct confidence, and the 5-section response format (links, steps, confidence sections all present).'),
    ('TestFeedbackEndpoint (4 tests): ', 'Confirms feedback submission for all three rating values (Good, Average, Poor) and rejects invalid ratings with HTTP 422.'),
    ('TestErrorHandling (2 tests): ', 'Tests the 503 response when Ollama is unreachable (via monkeypatching) and validates the frontend response format structure.'),
    ('TestConfigEndpoint (1 test): ', 'Verifies the root endpoint serves the SPA HTML correctly with proper content-type header.'),
]
for bold_part, text in test_classes:
    add_bullet(doc, text, bold_prefix=bold_part)

# 5.6
doc.add_heading('5.6 Task 6: Prompt Engineering', level=2)
doc.add_paragraph('The prompt engineering strategy evolved through two phases:')

doc.add_heading('Original Approach', level=3)
doc.add_paragraph(
    'Initially, the raw student question was sent directly to the LLM without any system '
    'instructions or context. This resulted in the model answering from its general training '
    'data, often hallucinating university policies, inventing fee amounts, and providing '
    'inconsistent advice.'
)
add_code_block(doc, '# Original prompt — raw question only\nprompt = question')

doc.add_heading('Improved Approach', level=3)
doc.add_paragraph(
    'The improved prompt uses a structured system instruction with constrained output rules, '
    'retrieved knowledge-base context, and conversation history. Critically, the 1B model '
    'struggled with complex formatting instructions, so we adopted a hybrid strategy:'
)
add_bullet(doc, 'The LLM prompt is kept minimal — just instruct the model to answer based on the provided UDSM information and never invent policies.', 'Strategy 1: ')
add_bullet(doc, 'All 5-section formatting is handled in Python code via _format_structured_response(), which uses regex action-verb extraction to derive steps from the natural language answer.', 'Strategy 2: ')
add_bullet(doc, 'The FAQ Direct Engine bypasses the LLM entirely for known questions, providing guaranteed accurate answers without any formatting concerns.', 'Strategy 3: ')

add_code_block(doc, '''# Improved prompt template
system = "You are the UDSM Student Support Assistant. Answer based only on the official UDSM information below. Never invent policies or fees."

if context:
    sections.append(f"Official UDSM information:\\n{context}")

sections.append(f"Student question: {question}")
prompt = "\\n\\n".join(sections)''')

# 5.7
doc.add_heading('5.7 Task 7: Error Handling', level=2)
doc.add_paragraph('The system handles four error situations as specified in the assignment:')

err_table = doc.add_table(rows=1, cols=3)
err_table.style = 'Light Grid Accent 1'
err_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = err_table.rows[0]
for i, text in enumerate(['Situation', 'Expected Behavior', 'Implementation']):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

errors = [
    ('Backend not running', 'Frontend shows connection error', 'try/catch block displays: "Network Error: Could not reach the API backend server."'),
    ('Model not running', 'Backend returns clear error', '/health returns status: degraded; /ask returns HTTP 503 "Local LLM service is currently unavailable"'),
    ('Empty question', 'Frontend asks user to enter a question', 'alert("Please enter a question.") + Pydantic min_length=1 validation'),
    ('Slow response', 'Frontend shows loading/spinner', 'Animated ping loader: "Retrieving database records..." during LLM generation'),
]
for sit, exp, impl in errors:
    add_table_row(err_table, [sit, exp, impl])

doc.add_page_break()

# 5.8
doc.add_heading('5.8 Task 8: Logging', level=2)
doc.add_paragraph('The backend implements comprehensive logging using Python\'s standard logging module:')

log_table = doc.add_table(rows=1, cols=3)
log_table.style = 'Light Grid Accent 1'
log_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = log_table.rows[0]
for i, text in enumerate(['Log File', 'Location', 'Content']):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

log_data = [
    ('app.log', 'backend/logs/app.log', 'Received questions, generated answers, RAG usage, category labels, confidence scores, timestamps, error details, system events'),
    ('feedback.log', 'backend/logs/feedback.log', 'User ratings (Good/Average/Poor) with associated question and answer text'),
]
for file, loc, content in log_data:
    add_table_row(log_table, [file, loc, content])

doc.add_paragraph()
doc.add_paragraph('Each log entry includes: timestamp, logger name, log level, and a structured message with interaction details. The log file rotates with each server restart to prevent unbounded growth.')

add_code_block(doc, '''# Sample log entry
2026-06-21 00:48:17,513 [INFO] llm_client - Sending request to Ollama (llama3.2:1b)
2026-06-21 00:49:04,538 [ERROR] llm_client - Failed to connect to Ollama: Read timed out.
2026-06-21 00:53:42,848 [INFO] student_support_backend - Health check endpoint called.''')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. TESTING AND RESULTS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Testing and Results', level=1)
doc.add_paragraph('Testing was executed in two phases: automated unit testing and manual error injection.')

doc.add_heading('Test Suite Results', level=2)
doc.add_paragraph('Running pytest tests/test_api.py -v produces the following results:')

test_results_table = doc.add_table(rows=1, cols=3)
test_results_table.style = 'Light Grid Accent 1'
test_results_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = test_results_table.rows[0]
for i, text in enumerate(['Test', 'Status', 'Description']):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

test_results = [
    ('test_health_returns_200', 'PASS', 'Health endpoint returns HTTP 200'),
    ('test_health_has_expected_fields', 'PASS', 'Response contains status, llm_connected, timestamp'),
    ('test_rejects_empty_question', 'PASS', 'Empty string returns 400'),
    ('test_rejects_whitespace_only', 'PASS', 'Whitespace returns 400'),
    ('test_rejects_missing_question', 'PASS', 'Missing field returns 400'),
    ('test_response_has_expected_fields', 'PASS', 'All 7 response fields present'),
    ('test_faq_direct_high_confidence', 'PASS', 'FAQ answers have High confidence'),
    ('test_response_includes_links', 'PASS', 'Useful Links section present'),
    ('test_response_includes_steps', 'PASS', 'Steps section present'),
    ('test_response_includes_confidence', 'PASS', 'Confidence section present'),
    ('test_feedback_saves', 'PASS', 'Feedback endpoint works'),
    ('test_feedback_rejects_invalid', 'PASS', 'Invalid rating returns 422'),
    ('test_feedback_accepts_average', 'PASS', 'Average rating accepted'),
    ('test_feedback_accepts_poor', 'PASS', 'Poor rating accepted'),
    ('test_returns_503_llm_down', 'PASS', '503 on LLM connection failure'),
    ('test_frontend_format', 'PASS', '5-section format confirmed'),
    ('test_spa_serves_index_html', 'PASS', 'SPA served at root with correct content-type'),
]
for name, status, desc in test_results:
    add_table_row(test_results_table, [name, status, desc])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. CHALLENGES ENCOUNTERED
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Challenges Encountered', level=1)

challenges = [
    ('Small Model Formatting Limitations', 
     'The Llama 3.2 1B model lacks the capacity to reliably follow complex formatting instructions. '
     'Initial attempts to include the 5-section format in the system prompt resulted in the model '
     'ignoring instructions, omitting sections, or producing malformed output. Resolution: we keep '
     'the LLM prompt minimal and handle all formatting in Python code via _format_structured_response() '
     'with regex-based step extraction from the natural language answer.'),
    ('Cold Start Timeouts', 
     'The first LLM generation after a cold start consistently exceeded the default 45-second timeout. '
     'This is because the model must be loaded into memory and the first inference is CPU-intensive. '
     'Resolution: increased the HTTP timeout to 120 seconds and added prompt truncation (max 4000 characters) '
     'to keep generation time within acceptable bounds.'),
    ('LLM Number Hallucination', 
     'The 1B model frequently invents fee amounts, dates, and policy numbers that sound plausible but '
     'are factually incorrect. Resolution: we inject retrieved knowledge-base context directly into the '
     'prompt as the authoritative source, use the FAQ Direct engine to bypass the LLM for known questions, '
     'and the accuracy check layer detects and flags suspicious answers.'),
    ('Windows Environment Permissions', 
     'Windows security policies occasionally blocked subprocess creation and file access. '
     'Resolution: explicit permission handling and virtual environment isolation.'),
]
for title, desc in challenges:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. BONUS FEATURES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Bonus Features', level=1)

doc.add_heading('Option B: Retrieval-Augmented Generation (RAG)', level=2)
doc.add_paragraph('Two-tier RAG system:')
add_bullet(doc, '11 knowledge-base markdown documents covering all UDSM service areas, loaded into memory at startup and indexed by keyword.', 'Knowledge Base: ')
add_bullet(doc, 'FAQ database (faq.json) with keyword-linked answers for 8 university service categories.', 'FAQ Database: ')
add_bullet(doc, 'Dual retrieval: keyword overlap scoring against KB documents + FAQ keyword matching.', 'Retrieval: ')
add_bullet(doc, 'FAQ Direct Engine bypasses LLM entirely when confidence >= 85% (sub-second responses).', 'Direct Mode: ')
add_bullet(doc, 'FAQ context injected into prompt when confidence is between 60% and 85%.', 'Context Injection: ')

doc.add_heading('Option E: Response Evaluation', level=2)
add_bullet(doc, 'User rating buttons (Good / Average / Poor) on each assistant message bubble.', 'UI: ')
add_bullet(doc, 'Submissions sent via POST /feedback endpoint.', 'API: ')
add_bullet(doc, 'All feedback logged to backend/logs/feedback.log with question, answer, and rating.', 'Logging: ')
add_bullet(doc, 'Visual confirmation displayed after rating submission.', 'Confirmation: ')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. INDUSTRY PRODUCTION REFLECTION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('9. Industry Production Reflection', level=1)

qa_pairs = [
    ('1. What are the main components of your deployed LLM system?',
     'The system comprises: (a) Ollama serving the local llama3.2:1b model via REST API on port 11434; '
     '(b) a FastAPI application providing four API endpoints with RAG, caching, conversation memory, '
     'and confidence scoring; (c) a Tailwind CSS SPA with responsive design and dark mode; (d) a '
     'knowledge base of 11 official UDSM markdown documents; (e) an FAQ database with keyword-matched '
     'answers; and (f) log files for interaction auditing and feedback collection. All components run '
     'locally on the same machine within a Python virtual environment.'),

    ('2. Why is FastAPI useful in this pipeline?',
     'FastAPI is highly asynchronous, which is crucial for handling multiple concurrent long-running LLM '
     'generation requests without blocking the server. It also automatically generates interactive Swagger '
     'documentation (/docs), validates request bodies using Pydantic schemas, provides structured HTTP '
     'exception handling, and offers high performance comparable to Node.js and Go.'),

    ('3. What role does your chosen LLM model play?',
     'llama3.2:1b acts as the natural language understanding and generation engine. It processes the '
     'student\'s question along with the retrieved FAQ/KB context, synthesizes the facts, and formulates '
     'a human-like, polite, and cohesive response. The 1-billion parameter size allows it to run on '
     'consumer hardware (CPU-only) while still producing reasonable responses for the constrained '
     'domain of university administrative queries.'),

    ('4. What role does the frontend play?',
     'The SAP serves as the interactive bridge between students and the backend. It captures user input, '
     'validates query length, displays loading spinners during LLM generation, renders formatted responses '
     'with confidence badges and source citations, captures user feedback ratings, and provides a '
     'responsive layout that works on mobile, tablet, and desktop devices. The Streamlit wrapper '
     'enables deployment on an alternative port while maintaining the same user experience.'),

    ('5. What is the difference between running the model locally and using an external API?',
     'Running locally provides complete data privacy (no student information leaves the server), zero '
     'cost per token, full offline availability, and no dependency on third-party services. However, '
     'it requires local hardware resources (RAM/CPU) and is limited to smaller, less capable models. '
     'External APIs like OpenAI\'s GPT-4 offer vastly more intelligence, require no local hardware, '
     'and update automatically. But they incur per-token costs, require constant internet connectivity, '
     'and raise data privacy concerns about sending student information to external servers.'),

    ('6. What security risks may exist if this system is deployed in an organisation?',
     'Key risks include: (a) Prompt injection — students craft inputs to bypass system rules and extract '
     'unauthorized information; (b) Denial of Service — rapid concurrent queries overwhelm the local '
     'CPU/GPU; (c) Log exposure — log files may inadvertently record sensitive student identifiers; '
     '(d) Model data leakage — the LLM might repeat confidential information from its training data; '
     '(e) Lack of authentication — anyone who can reach the server can use the API without restriction.'),

    ('7. What improvements would be needed before deploying this system in production?',
     'Critical improvements include: (a) Authentication — implement OAuth2 or SSO for student identity '
     'verification; (b) GPU acceleration — deploy on NVIDIA CUDA hardware for sub-second generation; '
     '(c) Vector database — upgrade keyword-based RAG to semantic embeddings (ChromaDB, Qdrant) for '
     'better context matching; (d) Containerization — package backend and frontend in Docker for reliable '
     'deployment and scaling; (e) Rate limiting — prevent DoS attacks via request throttling and '
     'per-user quotas; (f) HTTPS enforcement — encrypt all traffic between components.'),

    ('8. How would you monitor the system in real-world use?',
     'Monitoring would use: (a) Prometheus and Grafana — to track server hardware metrics (CPU/GPU '
     'utilization, memory usage) and FastAPI request latency/error rates; (b) LLM observability '
     'platforms like LangFuse or Phoenix — to track prompt token counts, generation latencies, '
     'and user rating trends; (c) Structured logging with correlation IDs — to trace individual '
     'requests through the entire pipeline; (d) Health check endpoints — to detect component failures '
     'and trigger automated alerts.'),

    ('9. How would you protect sensitive student information?',
     'Protection measures include: (a) Data sanitization — regex filters to strip personal identifiers '
     '(phone numbers, passwords, student IDs) from logs before storage; (b) HTTPS/TLS — encrypt all '
     'frontend-backend-LLM traffic; (c) Role-Based Access Control (RBAC) — restrict API access tokens '
     'to authenticated students only; (d) Audit trails — log all access attempts with timestamps and '
     'user IDs; (e) Data minimization — the LLM only receives the student\'s question, never their '
     'personal information; (f) Log retention policies — automatically purge old logs after a defined period.'),

    ('10. What challenges did you face during implementation?',
     'The primary challenges were: (a) Small model constraints — the 1B parameter model struggles with '
     'formatting instructions and number accuracy, requiring a code-side formatting approach; (b) Timeout '
     'management — cold-start LLM generation exceeded default timeouts, requiring increases to 120s '
     'and prompt truncation; (c) Windows environment — Python path and permission issues required '
     'explicit handling in configuration; (d) Model hallucination — the small model frequently invents '
     'plausible-sounding but incorrect information, mitigated by the RAG pipeline and FAQ Direct engine.'),
]

for question, answer in qa_pairs:
    doc.add_heading(question, level=2)
    doc.add_paragraph(answer)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('10. Conclusion', level=1)
doc.add_paragraph(
    'A self-hosted LLM application pipeline provides a private, secure, and cost-effective method to '
    'automate student inquiries. Through FastAPI, a Tailwind CSS SPA, and Ollama, we constructed a '
    'robust application that safely guides model outputs using an offline RAG pipeline with caching, '
    'conversation memory, confidence scoring, and structured response formatting.'
)
doc.add_paragraph(
    'The system demonstrates all key aspects of a production AI pipeline: environment isolation, '
    'local model serving, API design, frontend development, automated testing, comprehensive logging, '
    'error handling, and documentation. The architecture provides a production-ready template for '
    'enterprise AI systems that require data sovereignty, low latency, and complete control over '
    'the technology stack.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX A: PROJECT STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('Appendix A: Project Structure', level=1)

add_code_block(doc, '''student-support-llm/
├── backend/
│   ├── main.py              # FastAPI server (4 endpoints)
│   ├── llm_client.py        # LLM client (RAG, cache, memory, confidence)
│   ├── config.py            # Environment configuration
│   ├── faq.json             # FAQ knowledge base
│   └── logs/
│       ├── app.log          # Interaction logs
│       └── feedback.log     # User ratings
├── frontend/
│   ├── index.html           # Tailwind CSS SPA (responsive, dark mode)
│   └── app.py               # Streamlit wrapper (46 lines)
├── knowledge-base/          # 11 UDSM markdown RAG documents
├── tests/
│   └── test_api.py          # Pytest test suite (17 tests)
├── docs/
│   ├── screenshots/         # Evidence screenshots
│   └── report.md            # Technical report source
├── requirements.txt         # Python dependencies
├── README.md                # Setup and operations guide
└── generate_report.py       # Word document generator''')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX B: PROMPT COMPARISON
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('Appendix B: Prompt Comparison (Task 6)', level=1)

doc.add_heading('Original Prompt', level=2)
add_code_block(doc, '''{question}''')

doc.add_paragraph('Sample Output (Original Prompt):')
doc.add_paragraph(
    'Student: "I want to cheat in my exam, what are the rules?"'
)
doc.add_paragraph(
    'LLM (Hallucinated): "Cheating is generally not recommended, but some schools allow open notes '
    'or cheat sheets depending on the professor. Here are tips on how to prepare..."'
)
doc.add_paragraph('(Potentially dangerous or incorrect advice — the model invents policy from general knowledge.)')

doc.add_paragraph()

doc.add_heading('Improved Prompt (Simplified)', level=2)
add_code_block(doc, '''You are the UDSM Student Support Assistant. Answer based only on the official UDSM information below. Never invent policies or fees.

Official UDSM information:
{retrieved KB context}

Student question: {question}''')

doc.add_paragraph('Sample Output (Improved Prompt):')
doc.add_paragraph(
    'Student: "I want to cheat in my exam, what are the rules?"'
)
doc.add_paragraph(
    'LLM (Controlled): "Possession of unauthorized materials or cheating is strictly treated as '
    'academic dishonesty under the university examination rules, leading to immediate disciplinary '
    'action, suspension, or expulsion. Please refer to the Dean of Students\' office for code of '
    'conduct details."'
)

doc.add_paragraph()

doc.add_heading('Effect of Improved Prompt', level=2)
doc.add_paragraph(
    'Before improvement, the model answered from its general training data, often hallucinating '
    'university policies and providing potentially harmful advice. After improvement with RAG context '
    'injection, the model answers exclusively from the provided official documents, with clear safety '
    'guardrails for inappropriate questions. Furthermore, the _format_structured_response() method '
    'wraps the LLM output in a consistent 5-section format (Answer, Steps, Useful Links, Sources, '
    'Confidence) that no longer depends on the model\'s ability to follow formatting instructions.'
)

doc.add_page_break()

# ── Save ────────────────────────────────────────────────────────────────────
output_path = r'D:\CTFs\AI\AI\docs\IS_365_Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
